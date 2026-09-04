"""Flask server for Dossier authentication and resume analysis."""

import os
import re
import sqlite3
from markupsafe import escape
from io import BytesIO
from functools import wraps

from flask import Flask, jsonify, request, session, send_file
try:
    from flask_mail import Mail, Message
except ImportError:
    Mail = None
    Message = None
from werkzeug.security import check_password_hash, generate_password_hash

from model import Resume, analyze_ats, analyze_resume


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATABASE = os.path.join(BASE_DIR, "dossier.db")
app = Flask(__name__, static_folder=None)
app.config["SECRET_KEY"] = os.environ.get("DOSSIER_SECRET_KEY", "change-this-development-secret")

# ── Mail config (set these environment variables before running) ──
app.config["MAIL_SERVER"]   = os.environ.get("MAIL_SERVER", "smtp.gmail.com")
app.config["MAIL_PORT"]     = int(os.environ.get("MAIL_PORT", "587"))
app.config["MAIL_USE_TLS"]  = True
app.config["MAIL_USERNAME"] = os.environ.get("MAIL_USERNAME", "")   # your sending address
app.config["MAIL_PASSWORD"] = os.environ.get("MAIL_PASSWORD", "")   # app password / SMTP password
app.config["MAIL_DEFAULT_SENDER"] = os.environ.get("MAIL_USERNAME", "")
COMPANY_EMAIL = os.environ.get("COMPANY_EMAIL", "")   # inbox that receives new-signup alerts
mail = Mail(app) if Mail else None


def database_connection():
    connection = sqlite3.connect(DATABASE)
    connection.row_factory = sqlite3.Row
    return connection


# ── Single admin credential (set env vars; fallback for local dev only) ──
ADMIN_EMAIL    = os.environ.get("ADMIN_EMAIL",    "admin@dossier.app")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "change-admin-password")


def initialize_database():
    with database_connection() as connection:
        connection.execute(
            """CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                email TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
            )"""
        )
        user_columns = {row["name"] for row in connection.execute("PRAGMA table_info(users)")}
        if "name" not in user_columns:
            connection.execute("ALTER TABLE users ADD COLUMN name TEXT NOT NULL DEFAULT ''")
        if "phone" not in user_columns:
            connection.execute("ALTER TABLE users ADD COLUMN phone TEXT NOT NULL DEFAULT ''")
        if "welcome_sent_at" not in user_columns:
            connection.execute("ALTER TABLE users ADD COLUMN welcome_sent_at TEXT")
        connection.execute(
            """CREATE TABLE IF NOT EXISTS resume_plans (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                plan_type TEXT NOT NULL DEFAULT 'free',
                status TEXT NOT NULL DEFAULT 'active',
                updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(user_id) REFERENCES users(id)
            )"""
        )
        connection.execute(
            """CREATE TABLE IF NOT EXISTS resume_usage (
                user_id INTEGER PRIMARY KEY,
                resume_count INTEGER NOT NULL DEFAULT 0,
                updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(user_id) REFERENCES users(id)
            )"""
        )
        connection.execute(
            """CREATE TABLE IF NOT EXISTS contact_messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                name TEXT NOT NULL,
                email TEXT NOT NULL,
                message TEXT NOT NULL,
                reply TEXT,
                replied_at TEXT,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(user_id) REFERENCES users(id)
            )"""
        )


initialize_database()


def user_resume_access():
    if session.get("is_admin"):
        return {"isAdmin": True, "plan": "admin", "resumeCount": 0, "canCreate": True}
    user_id = session.get("user_id")
    if not user_id:
        return None
    with database_connection() as connection:
        plan = connection.execute(
            "SELECT plan_type, status FROM resume_plans WHERE user_id = ?", (user_id,)
        ).fetchone()
        usage = connection.execute(
            "SELECT resume_count FROM resume_usage WHERE user_id = ?", (user_id,)
        ).fetchone()
    plan_type = plan["plan_type"] if plan else "free"
    active_plan = plan and plan["status"] == "active"
    count = usage["resume_count"] if usage else 0
    unlimited = plan_type in {"pro", "pro-plus"} and active_plan
    return {
        "isAdmin": False,
        "plan": plan_type,
        "resumeCount": count,
        "canCreate": unlimited or count < 1,
    }


def current_user_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if "user_id" not in session:
            return jsonify(error="Please sign in first."), 401
        return view(*args, **kwargs)
    return wrapped


def admin_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not session.get("is_admin"):
            return jsonify(error="Admin access required."), 403
        return view(*args, **kwargs)
    return wrapped


def document_text(data, kind):
    if kind == "resume":
        resume = Resume.from_dict(data)
        sections = [resume.contact.get("name", ""), resume.contact.get("role", ""), resume.contact.get("email", ""), resume.summary]
        for title, items in (("Experience", resume.experience), ("Projects", resume.projects), ("Education", resume.education)):
            sections.append(title)
            for item in items:
                sections.extend(str(value) for value in item.values())
        sections.extend(["Skills", ", ".join(resume.skills)])
        return [section for section in sections if section]
    if kind == "cover-letter":
        name    = data.get("name", "")
        job     = data.get("job_title", "")
        email   = data.get("email", "")
        phone   = data.get("phone", "")
        loc     = data.get("location", "")
        date    = data.get("date", "")
        mgr     = data.get("hiring_manager", "Hiring Manager")
        company = data.get("company", "")
        addr    = data.get("company_address", "")
        opening = data.get("opening", "")
        body    = data.get("body", "")
        closing = data.get("closing", "")
        signoff = data.get("signoff", "Sincerely,")
        contact_line = "  ·  ".join(p for p in [email, phone, loc] if p)
        recipient = "  ".join(p for p in [mgr, company, addr] if p)
        return [name, contact_line, date, recipient, f"Re: {job}",
                f"Dear {mgr},", opening, body, closing, signoff, name]
    return [data.get("name", ""), data.get("role", ""), data.get("email", ""), data.get("location", ""), "Profile", data.get("profile", ""), "Experience", data.get("experience", ""), "Education", data.get("education", ""), "Skills", data.get("skills", ""), "Languages", data.get("languages", "")]


def export_document(data, kind, file_format):
    sections = document_text(data, kind)
    output = BytesIO()
    if file_format == "docx":
        from docx import Document
        document = Document()
        for index, section in enumerate(sections):
            document.add_heading(section, level=1 if index == 0 else 2) if section in {"Experience", "Projects", "Education", "Skills", "Profile", "Languages"} or index == 0 else document.add_paragraph(section)
        document.save(output)
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen.canvas import Canvas
        canvas = Canvas(output, pagesize=letter)
        y = 750
        for section in sections:
            for line in str(section).splitlines() or [""]:
                if y < 45:
                    canvas.showPage()
                    y = 750
                canvas.drawString(54, y, line[:110])
                y -= 17
            y -= 5
        canvas.save()
        mimetype = "application/pdf"
    output.seek(0)
    return output, mimetype


def request_credentials():
    data = request.get_json(silent=True) or {}
    email = str(data.get("email", "")).strip().lower()
    password = str(data.get("password", ""))
    if not re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", email):
        return None, None, ("Enter a valid email address.", 400)
    if len(password) < 8:
        return None, None, ("Password must be at least 8 characters.", 400)
    return email, password, None


def page(filename):
    return send_file(os.path.join(BASE_DIR, filename))


@app.get("/")
def index():
    return page("index.html")


@app.get("/about")
def about():
    return page("about.html")


@app.get("/contact")
def contact():
    return page("contact.html")


@app.get("/career-blog")
def career_blog():
    return page("career-blog.html")


@app.get("/cv-builder")
def cv_builder():
    if "user_id" not in session:
        from flask import redirect, url_for
        return redirect("/login")
    return page("cv-builder.html")


@app.get("/login")
def login_page():
    return page("login.html")


@app.get("/plans")
def plans_page():
    return page("plans.html")


@app.get("/templates")
def templates_page():
    return page("templates.html")


@app.get("/how-it-works")
def how_it_works_page():
    return page("how-it-works.html")


@app.get("/signup")
def signup_page():
    return page("signup.html")


@app.get("/cover-letter")
def cover_letter():
    if "user_id" not in session:
        from flask import redirect
        return redirect("/login")
    return page("cover-letter.html")


@app.get("/builder")
def builder():
    if "user_id" not in session and not session.get("is_admin"):
        from flask import redirect
        return redirect("/login")
    return page("html.html")


@app.get("/static/<path:filename>")
def static_file(filename):
    return send_file(os.path.join(BASE_DIR, filename))


WELCOME_HTML = """\
<!DOCTYPE html>
<html lang="en">
<head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<style>
  body{{margin:0;padding:0;background:#f4f6f8;font-family:'Segoe UI',Arial,sans-serif;color:#27364b}}
  .wrap{{max-width:600px;margin:40px auto;background:#fff;border-radius:12px;overflow:hidden;box-shadow:0 4px 18px #26364b14}}
  .header{{background:linear-gradient(135deg,#1a0a2e 0%,#3b1f6e 60%,#6d28d9 100%);padding:36px 40px 28px;text-align:center}}
  .header h1{{margin:0;color:#e9d5ff;font-size:28px;letter-spacing:-.02em}}
  .header small{{color:#c4b5fd;font-size:13px;letter-spacing:.06em}}
  .body{{padding:36px 40px}}
  .body p{{line-height:1.7;margin:0 0 18px;font-size:15px;color:#374151}}
  .cta{{text-align:center;margin:30px 0}}
  .cta a{{background:#7c3aed;color:#fff;text-decoration:none;padding:14px 34px;border-radius:30px;font-size:15px;font-weight:600;display:inline-block}}
  .sign{{margin-top:30px;font-size:15px;color:#374151}}
  .footer{{background:#f9fafb;border-top:1px solid #e5e7eb;padding:22px 40px;font-size:11px;color:#9ca3af;text-align:center;line-height:1.8}}
  .footer a{{color:#7c3aed;text-decoration:none}}
</style>
</head>
<body>
<div class="wrap">
  <div class="header">
    <h1>Dossier</h1>
    <small>resume workshop</small>
  </div>
  <div class="body">
    <p>Dear Dossier Job Seeker,</p>
    <p>Welcome to <strong>Dossier</strong>. Like you, we take your job search very seriously, offering the fastest way to build polished resumes and match them against real job descriptions.</p>
    <p>With your account you will have an all-access pass to our resume builder, CV builder, ATS scanner, and career blog — everything you need to job search with confidence.</p>
    <div class="cta"><a href="https://dossier.app/builder">Open Your Workspace Now</a></div>
    <p>Learn more <a href="https://dossier.app/about">About Us</a>, or contact Dossier at <a href="mailto:support@dossier.app">support@dossier.app</a>. We are here to help with any questions.</p>
    <p class="sign">Wishing You Great Success,<br><strong>The Dossier Team</strong></p>
  </div>
  <div class="footer">
    &copy; 2026, Dossier. All rights reserved.<br>
    Any questions? <a href="mailto:support@dossier.app">Contact us.</a><br><br>
    You&#39;ll always have access to your saved documents as long as your account remains active.
    For full details, please review our
    <a href="https://dossier.app/privacy">Privacy Policy</a> and
    <a href="https://dossier.app/terms">Terms &amp; Conditions</a>.
  </div>
</div>
</body></html>
"""


def send_welcome_email(user_name: str, user_email: str) -> bool:
    """Send one personalized welcome email and report whether delivery was attempted successfully."""
    if mail is None or not app.config["MAIL_USERNAME"] or not app.config["MAIL_PASSWORD"]:
        return False  # email not configured — skip silently
    try:
        welcome = Message(
            subject="Welcome to Dossier — your resume workspace is ready",
            recipients=[user_email],
            html=WELCOME_HTML.replace("Dear Dossier Job Seeker", f"Dear {escape(user_name)}"),
        )
        mail.send(welcome)
    except Exception:
        return False  # never block authentication if mail fails
    return True


def send_signup_emails(user_name: str, user_email: str) -> None:
    """Send the welcome email and a signup alert to the company inbox."""
    if send_welcome_email(user_name, user_email):
        with database_connection() as connection:
            connection.execute(
                "UPDATE users SET welcome_sent_at=CURRENT_TIMESTAMP WHERE email=?",
                (user_email,),
            )
    if COMPANY_EMAIL:
        try:
            alert = Message(
                subject=f"New signup: {user_email}",
                recipients=[COMPANY_EMAIL],
                body=f"A new user just signed up on Dossier.\n\nEmail: {user_email}\n",
            )
            mail.send(alert)
        except Exception:
            pass


@app.post("/api/signup")
def signup():
    data = request.get_json(silent=True) or {}
    name = " ".join(str(data.get("name", "")).strip().split())
    phone = str(data.get("phone", "")).strip()
    if len(name) < 2:
        return jsonify(error="Enter your full name."), 400
    if not re.fullmatch(r"[+\d()\-\s]{7,20}", phone):
        return jsonify(error="Enter a valid phone number."), 400
    email, password, error = request_credentials()
    if error:
        return jsonify(error=error[0]), error[1]
    try:
        with database_connection() as connection:
            cursor = connection.execute(
                "INSERT INTO users (name, phone, email, password_hash) VALUES (?, ?, ?, ?)",
                (name, phone, email, generate_password_hash(password)),
            )
            user_id = cursor.lastrowid
            session["user_id"] = user_id
            connection.execute(
                "INSERT INTO resume_plans (user_id) VALUES (?)", (user_id,)
            )
    except sqlite3.IntegrityError:
        return jsonify(error="An account with that email already exists."), 409
    send_signup_emails(name, email)
    return jsonify(user={"name": name, "phone": phone, "email": email}), 201


@app.post("/api/login")
def login():
    email, password, error = request_credentials()
    if error:
        return jsonify(error=error[0]), error[1]
    with database_connection() as connection:
        user = connection.execute("SELECT * FROM users WHERE email = ?", (email,)).fetchone()
    if user is None or not check_password_hash(user["password_hash"], password):
        return jsonify(error="Email or password is incorrect."), 401
    session["user_id"] = user["id"]
    if not user["welcome_sent_at"] and send_welcome_email(user["name"], user["email"]):
        with database_connection() as connection:
            connection.execute(
                "UPDATE users SET welcome_sent_at=CURRENT_TIMESTAMP WHERE id=?",
                (user["id"],),
            )
    return jsonify(user={"name": user["name"], "phone": user["phone"], "email": user["email"]}, resumeAccess=user_resume_access())


@app.get("/api/me")
def me():
    if session.get("is_admin"):
        return jsonify(user={"email": ADMIN_EMAIL, "isAdmin": True}, resumeAccess=user_resume_access())
    if "user_id" not in session:
        return jsonify(user=None)
    with database_connection() as connection:
        user = connection.execute("SELECT name, phone, email FROM users WHERE id = ?", (session["user_id"],)).fetchone()
    return jsonify(user=dict(user) if user else None, resumeAccess=user_resume_access() if user else None)


@app.get("/api/resume/access")
def resume_access():
    access = user_resume_access()
    if access is None:
        return jsonify(error="Please sign in first."), 401
    return jsonify(access=access)


@app.post("/api/resume/claim")
def claim_resume():
    access = user_resume_access()
    if access is None:
        return jsonify(error="Please sign in first."), 401
    if access["isAdmin"] or access["plan"] in {"pro", "pro-plus"}:
        return jsonify(ok=True, access=access)
    user_id = session["user_id"]
    with database_connection() as connection:
        connection.execute(
            "INSERT OR IGNORE INTO resume_usage (user_id, resume_count) VALUES (?, 0)",
            (user_id,),
        )
        cursor = connection.execute(
            "UPDATE resume_usage SET resume_count = resume_count + 1, updated_at=CURRENT_TIMESTAMP "
            "WHERE user_id = ? AND resume_count < 1",
            (user_id,),
        )
    if cursor.rowcount != 1:
        return jsonify(error="Your free resume is already used. Upgrade to Pro to create unlimited resumes.", upgrade=True), 402
    return jsonify(ok=True, access=user_resume_access())


@app.post("/api/logout")
def logout():
    session.clear()
    return jsonify(ok=True)


@app.post("/api/contact")
def contact_message():
    data = request.get_json(silent=True) or {}
    name = " ".join(str(data.get("name", "")).strip().split())
    email = str(data.get("email", "")).strip().lower()
    message = str(data.get("message", "")).strip()
    if len(name) < 2 or not re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", email) or len(message) < 3:
        return jsonify(error="Enter your name, a valid email, and a message."), 400
    with database_connection() as connection:
        connection.execute(
            "INSERT INTO contact_messages (user_id, name, email, message) VALUES (?, ?, ?, ?)",
            (session.get("user_id"), name, email, message),
        )
    return jsonify(ok=True), 201


@app.post("/api/analyze")
@current_user_required
def analyze():
    data = request.get_json(silent=True) or {}
    resume_text = str(data.get("resume_text", "")).strip()
    job_description = str(data.get("job_description", "")).strip()
    if not resume_text or not job_description:
        return jsonify(error="Resume text and job description are required."), 400
    result = analyze_resume(resume_text, job_description)
    return jsonify(result)


@app.post("/api/analyze-upload")
@current_user_required
def analyze_upload():
    uploaded = request.files.get("resume")
    job_description = request.form.get("job_description", "").strip()
    if uploaded is None or not uploaded.filename:
        return jsonify(error="Choose a resume file first."), 400
    if not job_description:
        return jsonify(error="Job description is required."), 400
    extension = os.path.splitext(uploaded.filename)[1].lower()
    try:
        if extension == ".txt":
            resume_text = uploaded.read().decode("utf-8", errors="ignore")
        elif extension == ".pdf":
            from pypdf import PdfReader
            resume_text = "\n".join(page.extract_text() or "" for page in PdfReader(uploaded).pages)
        elif extension == ".docx":
            from docx import Document
            resume_text = "\n".join(paragraph.text for paragraph in Document(uploaded).paragraphs)
        else:
            return jsonify(error="Supported files are PDF, DOCX, and TXT."), 415
    except Exception:
        return jsonify(error="The file could not be read. Try exporting it as PDF or DOCX."), 400
    if not resume_text.strip():
        return jsonify(error="No readable text was found in that file."), 400
    return jsonify(resumeText=resume_text, result=analyze_resume(resume_text, job_description))


@app.post("/api/extract-resume")
@current_user_required
def extract_resume():
    uploaded = request.files.get("resume")
    if uploaded is None or not uploaded.filename:
        return jsonify(error="Choose a resume file first."), 400
    extension = os.path.splitext(uploaded.filename)[1].lower()
    try:
        if extension == ".txt":
            resume_text = uploaded.read().decode("utf-8", errors="ignore")
        elif extension == ".pdf":
            from pypdf import PdfReader
            resume_text = "\n".join(page.extract_text() or "" for page in PdfReader(uploaded).pages)
        elif extension == ".docx":
            from docx import Document
            resume_text = "\n".join(paragraph.text for paragraph in Document(uploaded).paragraphs)
        else:
            return jsonify(error="Supported files are PDF, DOCX, and TXT."), 415
    except Exception:
        return jsonify(error="The file could not be read. Try exporting it as PDF or DOCX."), 400
    if not resume_text.strip():
        return jsonify(error="No readable text was found in that file."), 400
    return jsonify(resumeText=resume_text)


@app.post("/api/download/<kind>/<file_format>")
@current_user_required
def download_document(kind, file_format):
    if kind not in {"resume", "cv", "cover-letter"} or file_format not in {"docx", "pdf"}:
        return jsonify(error="Unsupported export format."), 400
    data = request.get_json(silent=True) or {}
    output, mimetype = export_document(data, kind, file_format)
    return send_file(output, mimetype=mimetype, as_attachment=True, download_name=f"dossier-{kind}.{file_format}")


@app.post("/api/ats-analyze")
@current_user_required
def ats_analyze():
    data = request.get_json(silent=True) or {}
    job_description = str(data.get("job_description", "")).strip()
    resume_data = data.get("resume") or {}
    if not job_description or not resume_data:
        return jsonify(error="Resume and job description are required."), 400
    return jsonify(analyze_ats(Resume.from_dict(resume_data), job_description))


@app.get("/admin/login")
@app.get("/admin-login")
def admin_login_page():
    if session.get("is_admin"):
        from flask import redirect
        return redirect("/admin")
    return page("admin-login.html")


@app.post("/api/admin/login")
def admin_login():
    data = request.get_json(silent=True) or {}
    email = str(data.get("email", "")).strip().lower()
    password = str(data.get("password", ""))
    if email != ADMIN_EMAIL.lower() or password != ADMIN_PASSWORD:
        return jsonify(error="Invalid admin credentials."), 401
    session.clear()
    session["is_admin"] = True
    return jsonify(ok=True)


@app.post("/api/admin/logout")
def admin_logout():
    session.clear()
    return jsonify(ok=True)


@app.get("/admin")
def admin_page():
    if not session.get("is_admin"):
        from flask import redirect
        return redirect("/admin/login")
    return page("admin.html")


@app.get("/api/admin/stats")
@admin_required
def admin_stats():
    with database_connection() as connection:
        total_users = connection.execute("SELECT COUNT(*) as c FROM users").fetchone()["c"]
        new_today = connection.execute(
            "SELECT COUNT(*) as c FROM users WHERE DATE(created_at) = DATE('now')"
        ).fetchone()["c"]
        new_week = connection.execute(
            "SELECT COUNT(*) as c FROM users WHERE created_at >= DATE('now', '-7 days')"
        ).fetchone()["c"]
        signups_by_day = connection.execute(
            """SELECT DATE(created_at) as day, COUNT(*) as count
               FROM users GROUP BY day ORDER BY day DESC LIMIT 30"""
        ).fetchall()
        plan_breakdown = connection.execute(
            """SELECT plan_type, status, COUNT(*) as count
               FROM resume_plans GROUP BY plan_type, status"""
        ).fetchall()
        recent_users = connection.execute(
            """SELECT u.email, u.created_at, COALESCE(rp.plan_type,'free') as plan,
               COALESCE(rp.status,'active') as status
               FROM users u LEFT JOIN resume_plans rp ON rp.user_id = u.id
               ORDER BY u.created_at DESC LIMIT 50"""
        ).fetchall()
        contact_messages = connection.execute(
            """SELECT id, name, email, message, reply, replied_at, created_at
               FROM contact_messages ORDER BY created_at DESC LIMIT 100"""
        ).fetchall()
    return jsonify(
        totalUsers=total_users,
        newToday=new_today,
        newThisWeek=new_week,
        signupsByDay=[dict(r) for r in signups_by_day],
        planBreakdown=[dict(r) for r in plan_breakdown],
        recentUsers=[dict(r) for r in recent_users],
        contactMessages=[dict(r) for r in contact_messages],
    )


@app.post("/api/admin/contact/<int:message_id>/reply")
@admin_required
def reply_to_contact(message_id):
    data = request.get_json(silent=True) or {}
    reply = str(data.get("reply", "")).strip()
    if not reply:
        return jsonify(error="Write a reply first."), 400
    with database_connection() as connection:
        message = connection.execute(
            "SELECT email, name FROM contact_messages WHERE id = ?", (message_id,)
        ).fetchone()
    if message is None:
        return jsonify(error="Message not found."), 404
    if mail is None or not app.config["MAIL_USERNAME"] or not app.config["MAIL_PASSWORD"]:
        return jsonify(error="Email delivery is not configured. Set MAIL_USERNAME and MAIL_PASSWORD, then restart the app."), 503
    try:
        mail.send(Message(
            subject="Reply from Dossier Customer Care",
            recipients=[message["email"]],
            body=f"Hello {message['name']},\n\n{reply}\n\nRegards,\nDossier Customer Care",
        ))
    except Exception:
        return jsonify(error="The reply could not be sent."), 502
    with database_connection() as connection:
        connection.execute(
            "UPDATE contact_messages SET reply=?, replied_at=CURRENT_TIMESTAMP WHERE id=?",
            (reply, message_id),
        )
    return jsonify(ok=True)


@app.post("/api/admin/user/<int:user_id>/plan")
@admin_required
def update_user_plan(user_id):
    data = request.get_json(silent=True) or {}
    plan_type = data.get("plan_type", "free")
    status = data.get("status", "active")
    with database_connection() as connection:
        existing = connection.execute("SELECT id FROM resume_plans WHERE user_id = ?", (user_id,)).fetchone()
        if existing:
            connection.execute(
                "UPDATE resume_plans SET plan_type=?, status=?, updated_at=CURRENT_TIMESTAMP WHERE user_id=?",
                (plan_type, status, user_id)
            )
        else:
            connection.execute(
                "INSERT INTO resume_plans (user_id, plan_type, status) VALUES (?,?,?)",
                (user_id, plan_type, status)
            )
    return jsonify(ok=True)


if __name__ == "__main__":
    app.run(
        host="0.0.0.0",
        port=int(os.environ.get("PORT", 5000))
    )